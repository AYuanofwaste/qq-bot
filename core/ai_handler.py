"""OpenAI-powered handler — 聊天优先，工具辅助。支持多轮对话记忆。"""

import asyncio
import json
import logging
import os
import re
import time
from pathlib import Path

from openai import AsyncOpenAI

from mcp_client.client import McpManager

logger = logging.getLogger("qq-bot.ai")


def _create_pdf(path_str: str) -> str:
    import img2pdf
    src = Path(path_str).resolve()
    images = []
    if src.is_dir():
        for f in sorted(src.iterdir()):
            if f.is_file() and f.suffix.lower() in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
                images.append(str(f))
    elif src.is_file():
        images.append(str(src))
    if not images:
        raise RuntimeError(f"没有找到可转换的图片: {path_str}")
    pdf_path = str(src.with_name(src.name + ".pdf"))
    with open(pdf_path, "wb") as f:
        f.write(img2pdf.convert(images))
    logger.info("PDF 已创建: %s (%d 页)", pdf_path, len(images))
    if src.is_dir():
        import shutil
        shutil.rmtree(src, ignore_errors=True)
    elif src.is_file():
        src.unlink(missing_ok=True)
    return pdf_path


def _create_docx(path_str: str) -> str:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    src = Path(path_str).resolve()
    doc = Document()
    if src.is_file() and src.suffix.lower() == ".gif":
        doc.add_picture(str(src), width=Inches(5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif src.is_dir():
        images = sorted([f for f in src.iterdir()
                         if f.is_file() and f.suffix.lower() in (".jpg", ".jpeg", ".png", ".gif", ".webp")])
        for img in images:
            doc.add_picture(str(img), width=Inches(5))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docx_path = str(src.with_name(src.name + ".docx"))
    doc.save(docx_path)
    logger.info("DOCX 已创建: %s", docx_path)
    if src.is_dir():
        import shutil
        shutil.rmtree(src, ignore_errors=True)
    elif src.is_file():
        src.unlink(missing_ok=True)
    return docx_path

SYSTEM_PROMPT = ""


def set_system_prompt(prompt: str):
    global SYSTEM_PROMPT
    SYSTEM_PROMPT = prompt


def _build_tools():
    return [
        {
            "type": "function",
            "function": {
                "name": "download_jm_comic",
                "description": "从禁漫天堂下载漫画并转换为PDF文件（参数 jm_id 是纯数字ID，不是链接）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "jm_id": {
                            "type": "string",
                            "description": "JM漫画ID，例如 123 或 350234"
                        },
                        "output_dir": {
                            "type": "string",
                            "description": "PDF文件保存目录（可选，默认临时目录）"
                        }
                    },
                    "required": ["jm_id"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "search_jm_comic",
                "description": "搜索禁漫天堂漫画（同官网搜索语法：+标签=必须含该标签，-标签=排除该标签）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "keyword": {
                            "type": "string",
                            "description": "搜索关键词，多个标签用空格分隔"
                        },
                        "page": {
                            "type": "integer",
                            "description": "页码（默认1）"
                        }
                    },
                    "required": ["keyword"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "search_illust",
                "description": "搜索Pixiv插画，根据关键词查找作品（参数 word 是文字关键词，不是数字ID）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "word": {"type": "string", "description": "搜索关键词或标签"},
                        "search_target": {
                            "type": "string",
                            "enum": ["partial_match_for_tags", "exact_match_for_tags",
                                     "title_and_caption", "keyword"],
                            "description": "搜索目标类型"
                        },
                        "sort": {
                            "type": "string",
                            "enum": ["date_desc", "date_asc", "popular_desc", "popular_asc"],
                            "description": "排序方式"
                        },
                        "duration": {
                            "type": "string",
                            "enum": ["within_last_day", "within_last_week", "within_last_month"],
                            "description": "时间范围"
                        }
                    },
                    "required": ["word"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "illust_detail",
                "description": "获取Pixiv插画的详细信息，包括标题、作者、标签、尺寸等（参数 illust_id 是纯数字ID）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "illust_id": {"type": "integer", "description": "插画ID"}
                    },
                    "required": ["illust_id"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "illust_ranking",
                "description": "获取Pixiv插画排行榜（按 mode 和 date 参数查询）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "mode": {
                            "type": "string",
                            "enum": ["day", "week", "month", "day_male", "day_female",
                                     "week_original", "week_rookie"],
                            "description": "排行榜模式（日/周/月/男性/女性等）"
                        },
                        "date": {"type": "string", "description": "指定日期，格式 YYYY-MM-DD"}
                    }
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "trending_tags_illust",
                "description": "获取Pixiv当前热门搜索标签（无需参数）",
                "parameters": {"type": "object", "properties": {}}
            }
        },
        {
            "type": "function",
            "function": {
                "name": "illust_recommended",
                "description": "获取Pixiv推荐插画（无需参数）",
                "parameters": {"type": "object", "properties": {}}
            }
        },
        {
            "type": "function",
            "function": {
                "name": "illust_related",
                "description": "获取与指定插画相关的作品推荐（参数 illust_id 是纯数字ID）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "illust_id": {"type": "integer", "description": "源插画ID"}
                    },
                    "required": ["illust_id"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "search_user",
                "description": "搜索Pixiv用户（参数 word 是用户名或关键词）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "word": {"type": "string", "description": "用户名或关键词"}
                    },
                    "required": ["word"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "download_pixiv_illust",
                "description": "下载Pixiv插画到本地（参数 illust_id 是纯数字ID，不是链接）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "illust_id": {"type": "integer", "description": "插画ID"}
                    },
                    "required": ["illust_id"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "search_bilibili_live",
                "description": "在B站搜索正在直播的主播（参数 keyword 是主播名/关键词，不是链接）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "keyword": {"type": "string", "description": "搜索关键词（主播名、游戏名等）"},
                        "page": {"type": "integer", "description": "页码（默认1）"}
                    },
                    "required": ["keyword"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "bilibili_live_following",
                "description": "查看你B站关注的正在直播的主播列表（无需参数）",
                "parameters": {"type": "object", "properties": {}}
            }
        },
        {
            "type": "function",
            "function": {
                "name": "download_bilibili_video",
                "description": "下载B站视频（参数 url 必须是 b23.tv/BV/bilibili.com 格式的链接，不是数字ID）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "url": {"type": "string", "description": "B站视频链接或BV号"}
                    },
                    "required": ["url"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "synthesize_tts",
                "description": "合成语音并发送语音消息（用户说「说一下xxx」「语音说xxx」「读一下xxx」等含朗读/语音含义时调用）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "text": {"type": "string", "description": "要朗读的文本"},
                        "language": {"type": "string", "enum": ["ZH","JP","EN"], "description": "语言（可选，默认当前语言）"}
                    },
                    "required": ["text"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "switch_tts_language",
                "description": "切换TTS默认语言（用户说「切换中文」「切换日文」「切换英文」时调用）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "language": {"type": "string", "enum": ["ZH","JP","EN"], "description": "目标语言"}
                    },
                    "required": ["language"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "toggle_ai_tts",
                "description": "开启或关闭AI自动语音功能（用户说「开启AI语音」「关闭AI语音」「打开自动语音」时调用，不传参数则切换）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "enabled": {"type": "boolean", "description": "true=开启 false=关闭，不传则切换"}
                    }
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "switch_persona",
                "description": "切换AI人设和TTS模型（用户说「切换东雪莲」「切换永雏塔菲」「切换到xxx」时调用）",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string", "description": "角色名，可用: Azuma / Taffy"}
                    },
                    "required": ["name"]
                }
            }
        },
    ]


_JM_OUTPUT = Path(__file__).resolve().parent.parent / "jmcomic" / "output"


class AIHandler:
    HISTORY_MAX_ROUNDS = 10
    HISTORY_TTL = 1800
    _histories: dict[str, dict] = {}

    def __init__(self, mcp: McpManager):
        self.mcp = mcp
        self.model = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("环境变量 OPENAI_API_KEY 未设置")

        kwargs = {"api_key": api_key}
        base_url = os.environ.get("OPENAI_BASE_URL")
        if base_url:
            kwargs["base_url"] = base_url

        self.client = AsyncOpenAI(**kwargs)
        self.tools = _build_tools()
        self.func_map = self._build_func_map()

    def _build_func_map(self):
        return {
            "download_jm_comic": self._call_jm_download,
            "search_jm_comic": self._call_jm_search,
            "search_illust": self._call_pixiv_search,
            "illust_detail": self._call_pixiv_detail,
            "illust_ranking": self._call_pixiv_ranking,
            "trending_tags_illust": self._call_pixiv_trending,
            "illust_recommended": self._call_pixiv_recommended,
            "illust_related": self._call_pixiv_related,
            "search_user": self._call_pixiv_search_user,
            "download_pixiv_illust": self._call_pixiv_download,
            "search_bilibili_live": self._call_bili_search_live,
            "bilibili_live_following": self._call_bili_live_following,
            "download_bilibili_video": self._call_bili_download,
            "synthesize_tts": self._call_synthesize_tts,
            "switch_tts_language": self._call_switch_tts_language,
            "toggle_ai_tts": self._call_toggle_ai_tts,
            "switch_persona": self._call_switch_persona,
        }

    def _get_history(self, user_id: str) -> list[dict]:
        sess = self._histories.get(user_id)
        if not sess:
            return []
        if time.time() - sess.get("ts", 0) > self.HISTORY_TTL:
            self._histories.pop(user_id, None)
            return []
        return sess["messages"]

    def _append_history(self, user_id: str, role: str, content: str):
        if user_id not in self._histories:
            self._histories[user_id] = {"messages": [], "ts": time.time()}
        sess = self._histories[user_id]
        sess["ts"] = time.time()
        sess["messages"].append({"role": role, "content": content})
        max_msgs = self.HISTORY_MAX_ROUNDS * 2
        if len(sess["messages"]) > max_msgs:
            sess["messages"] = sess["messages"][-max_msgs:]

    def clear_history(self, user_id: str):
        self._histories.pop(user_id, None)

    async def _call_jm_download(self, jm_id: str, output_dir: str = None):
        if not output_dir:
            output_dir = str(_JM_OUTPUT)
        pdf_path = await self.mcp.download_jm_comic(jm_id=jm_id, output_dir=output_dir)
        if isinstance(pdf_path, str) and pdf_path.endswith(".pdf"):
            self._task_files().append(pdf_path)
        return pdf_path

    async def _call_jm_search(self, keyword: str, page: int = 1):
        return await self.mcp.search_jm_comic(keyword=keyword, page=page)

    async def _call_pixiv_search(self, **kwargs):
        kwargs = {k: v for k, v in kwargs.items() if v is not None}
        return await self.mcp.search_illust(**kwargs)

    async def _call_pixiv_detail(self, illust_id: int):
        return await self.mcp.illust_detail(illust_id=illust_id)

    async def _call_pixiv_ranking(self, **kwargs):
        kwargs = {k: v for k, v in kwargs.items() if v is not None}
        return await self.mcp.illust_ranking(**kwargs)

    async def _call_pixiv_trending(self):
        return await self.mcp.trending_tags_illust()

    async def _call_pixiv_recommended(self):
        return await self.mcp.illust_recommended()

    async def _call_pixiv_related(self, illust_id: int):
        return await self.mcp.illust_related(illust_id=illust_id)

    async def _call_pixiv_search_user(self, word: str):
        return await self.mcp.search_user(word=word)

    async def _call_pixiv_download(self, illust_id: int):
        raw = await self.mcp.download(illust_id=illust_id, wait=True)
        data = json.loads(raw) if isinstance(raw, str) else raw
        if isinstance(data, dict):
            if data.get("error"):
                logger.error("Download %d failed: %s", illust_id, data["error"])
                raise RuntimeError(data["error"])
            if data.get("status") == "failed":
                err_msg = data.get("error", "下载失败")
                logger.error("Download %d failed: %s", illust_id, err_msg)
                raise RuntimeError(err_msg)
            rp = data.get("result_path")
            if rp:
                p = Path(rp)
                if p.is_file() and p.suffix.lower() == ".gif":
                    docx_path = await self._run_sync(_create_docx, str(p))
                    self._task_files().append(docx_path)
                else:
                    pdf_path = await self._run_sync(_create_pdf, str(p))
                    self._task_files().append(pdf_path)
        return raw

    async def _call_bili_search_live(self, keyword: str, page: int = 1):
        return await self.mcp.bili_search_live(keyword=keyword, page=page)

    async def _call_bili_live_following(self):
        return await self.mcp.bili_live_following()

    async def _call_bili_download(self, url: str):
        result = await self.mcp.bili_download_video(url=url)
        path = result.strip()
        if path and Path(path).is_file():
            self._task_files().append(path)
        return result

    async def _call_synthesize_tts(self, text: str, language: str = None):
        kwargs = {"text": text, "model_id": 0, "speaker_name": "Azuma", "language": language or "ZH"}
        from tts import tts_mgr
        kwargs["model_id"] = tts_mgr.model_id
        kwargs["speaker_name"] = tts_mgr.speaker_name
        wav_path = await self.mcp.synthesize(**kwargs)
        if wav_path.startswith("ERROR:"):
            return wav_path
        self._task_files().append(wav_path)
        return "语音已发送"

    async def _call_switch_tts_language(self, language: str):
        from tts.commands import handle_tts_switch
        handle_tts_switch._default_lang = language.upper()
        return f"已切换为{language.upper()}模式"

    async def _call_toggle_ai_tts(self, enabled: bool = None):
        from tts.commands import handle_tts_switch
        current = getattr(handle_tts_switch, "_ai_tts_enabled", True)
        if enabled is None:
            enabled = not current
        handle_tts_switch._ai_tts_enabled = enabled
        return f"AI 自动语音已{'开启' if enabled else '关闭'}"

    async def _call_switch_persona(self, name: str):
        from tts import tts_mgr
        from core.ai_handler import set_system_prompt
        prompt = tts_mgr.switch_persona(name)
        set_system_prompt(prompt)
        return f"已切换为{name}"

    _files_by_task: "dict[int, list[str]]" = {}

    @staticmethod
    async def _run_sync(func, *args, **kwargs):
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, lambda: func(*args, **kwargs))

    def _task_files(self) -> list[str]:
        try:
            task = asyncio.current_task()
        except RuntimeError:
            task = None
        key = id(task) if task else 0
        if key not in self._files_by_task:
            self._files_by_task[key] = []
        return self._files_by_task[key]

    async def handle(self, text: str, user_id: str = "",
                     on_tool_start=None) -> tuple[str, list[str]]:
        """
        on_tool_start: async callback(tool_name: str) called before each tool execution.
                       Bot can use this to send progress messages.
        """
        if not text or not text.strip():
            return "请输入你想说的话～", []

        task_files = self._task_files()
        task_files.clear()

        history = self._get_history(user_id) if user_id else []
        system_content = SYSTEM_PROMPT or "你是东雪莲，一个可爱的虚拟主播。关注东雪莲谢谢喵～"
        messages = [{"role": "system", "content": system_content}] + history + [
            {"role": "user", "content": text}
        ]

        try:
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                tools=self.tools,
                tool_choice="auto",
                temperature=0.8,
            )
        except Exception as e:
            logger.error("OpenAI first call failed: %s", e)
            return self._format_api_error(e), task_files[:]

        if not response.choices:
            return "AI 没有返回任何回复，请重试。", task_files[:]

        msg = response.choices[0].message

        if not msg.tool_calls:
            reply = msg.content or "好的，收到！"
            reply = re.sub(r'<tool_calls>.*?</tool_calls>', '', reply, flags=re.DOTALL).strip()
            if user_id:
                self._append_history(user_id, "user", text)
                self._append_history(user_id, "assistant", reply)
            return reply, task_files[:]

        assistant_msg = {"role": "assistant", "content": msg.content or None}
        if msg.tool_calls:
            assistant_msg["tool_calls"] = [
                {"id": tc.id, "type": "function",
                 "function": {"name": tc.function.name, "arguments": tc.function.arguments}}
                for tc in msg.tool_calls
            ]
        messages.append(assistant_msg)

        for tc in msg.tool_calls:
            func = self.func_map.get(tc.function.name)
            if not func:
                result = json.dumps({"error": f"未知函数: {tc.function.name}"})
            else:
                try:
                    args = json.loads(tc.function.arguments)
                    logger.info("Calling %s with %s", tc.function.name, args)
                    if on_tool_start:
                        await on_tool_start(tc.function.name)
                    raw = await func(**args)
                    if isinstance(raw, str):
                        result = raw
                    else:
                        result = json.dumps(raw, ensure_ascii=False)
                except Exception as e:
                    logger.error("Tool %s failed: %s", tc.function.name, e)
                    result = json.dumps({"error": str(e)})

            messages.append({
                "role": "tool",
                "tool_call_id": tc.id,
                "content": result,
            })

        try:
            final = await self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.8,
            )
        except Exception as e:
            logger.error("OpenAI second call failed: %s", e)
            return self._format_api_error(e), task_files[:]

        if not final.choices:
            return "处理完成，但 AI 未能生成回复，请重试。", task_files[:]

        reply = final.choices[0].message.content or "处理完成。"
        reply = re.sub(r'<tool_calls>.*?</tool_calls>', '', reply, flags=re.DOTALL).strip()
        if user_id:
            self._append_history(user_id, "user", text)
            self._append_history(user_id, "assistant", reply)
        return reply, task_files[:]

    @staticmethod
    def _format_api_error(e: Exception) -> str:
        msg = str(e)
        if "rate_limit" in msg.lower() or "429" in msg:
            return "⏳ AI 请求太频繁，请等几秒再试"
        if "timeout" in msg.lower() or "timed out" in msg.lower():
            return "⏳ AI 响应超时，请重试"
        if "insufficient_quota" in msg.lower() or "quota" in msg.lower():
            return "❌ AI 额度已用完，请联系管理员充值"
        if "auth" in msg.lower() or "api key" in msg.lower() or "401" in msg:
            return "❌ AI API Key 无效，请联系管理员检查配置"
        if "connection" in msg.lower() or "connect" in msg.lower():
            return "❌ 无法连接 AI 服务，请检查网络"
        return f"❌ AI 处理出错: {e}"
