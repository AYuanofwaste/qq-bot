"""QQ Bot — connects NapCatQQ (OneBot v11) with JMComic, Pixiv, Bilibili services via MCP."""

import asyncio
import json
import logging
import os
import re
import sys
import time
from pathlib import Path

from dotenv import load_dotenv
from aiocqhttp import CQHttp, Event
from aiocqhttp.message import Message

from core.ai_handler import AIHandler
from core.helpers import _split_args
from bilibili.dl import extract_bvid
from bilibili.live import format_live_page
from news.global_ import SOURCE_COMMANDS, fetch_global_news
from news.domestic import fetch_news
from mcp_client.client import McpManager


load_dotenv(Path(__file__).resolve().parent / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
    stream=sys.stderr,
)
logger = logging.getLogger("qq-bot")

BASE_DIR = Path(__file__).resolve().parent
JM_OUTPUT = BASE_DIR / "jmcomic" / "output"
BILI_OUTPUT = BASE_DIR / "bilibili" / "downloads"
PIXIV_OUTPUT = BASE_DIR / "pixiv" / "downloads"

_mcp_manager: McpManager | None = None


def _get_mcp() -> McpManager:
    global _mcp_manager
    if _mcp_manager is None:
        raise RuntimeError("MCP manager not initialized")
    return _mcp_manager


_ai_handler: AIHandler | None = None


def _get_ai() -> AIHandler:
    global _ai_handler
    if _ai_handler is None:
        _ai_handler = AIHandler(_get_mcp())
    return _ai_handler


_token = os.environ.get("BOT_TOKEN", "")
bot = CQHttp(access_token=_token) if _token else CQHttp()


# ── Helpers ──────────────────────────────────────────────────────

async def _send_reply(event: Event, message: str):
    await bot.send(event, Message(message))


MAX_QQ_FILE_SIZE = 50 * 1024 * 1024


async def _send_file(event: Event, file_path: str):
    path = Path(file_path).resolve()
    if not path.exists():
        logger.warning("文件不存在，无法发送: %s", path)
        return
    size_mb = path.stat().st_size / 1024 / 1024
    if size_mb > 50:
        logger.warning("文件过大 (%.1fMB)，跳过发送: %s", size_mb, path.name)
        return
    file_abs = path.as_posix()
    file_name = path.name
    try:
        cq = f"[CQ:file,file=file:///{file_abs},name={file_name}]"
        await bot.send(event, Message(cq))
        logger.info("文件已发送: %s (%.1fMB)", file_name, size_mb)
    except Exception as e:
        logger.warning("CQ:file 发送失败，尝试 upload: %s", e)
        if event.message_type == "private":
            await bot.call_action("upload_private_file", user_id=event.user_id, file=str(path), name=file_name)
        else:
            await bot.call_action("upload_group_file", group_id=event.group_id, file=str(path), name=file_name)
        logger.info("文件已发送 (upload): %s", file_name)


async def _send_video(event: Event, file_path: str):
    path = Path(file_path).resolve()
    if not path.exists():
        logger.warning("视频不存在: %s", path)
        return
    try:
        file_abs = path.as_posix()
        cq = f"[CQ:video,file=file:///{file_abs}]"
        await bot.send(event, Message(cq))
        logger.info("视频已发送: %s", path)
    except Exception as e:
        logger.warning("CQ:video 发送失败，改用文件发送: %s", e)
        await _send_file(event, file_path)


def _create_pdf(path_str: str) -> str:
    import img2pdf
    src = Path(path_str).resolve()
    images = []
    if src.is_dir():
        for f in sorted(src.iterdir()):
            if f.is_file() and f.suffix.lower() in (".jpg", ".jpeg", ".png", ".gif", ".webp"):
                images.append(str(f))
    elif src.is_file():
        images.append(str(src))
    if not images:
        raise RuntimeError(f"没有找到可转换的图片: {path_str}")
    pdf_path = str(src.with_name(src.name + ".pdf"))
    with open(pdf_path, "wb") as f:
        f.write(img2pdf.convert(images))
    logger.info("PDF 已创建: %s (%d 页)", pdf_path, len(images))
    if src.is_dir():
        import shutil
        shutil.rmtree(src, ignore_errors=True)
    elif src.is_file():
        src.unlink(missing_ok=True)
    return pdf_path


def _create_docx(path_str: str) -> str:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    src = Path(path_str).resolve()
    doc = Document()
    if src.is_file() and src.suffix.lower() == ".gif":
        doc.add_picture(str(src), width=Inches(5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif src.is_dir():
        images = sorted([f for f in src.iterdir()
                         if f.is_file() and f.suffix.lower() in (".jpg", ".jpeg", ".png", ".gif", ".webp")])
        for img in images:
            doc.add_picture(str(img), width=Inches(5))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docx_path = str(src.with_name(src.name + ".docx"))
    doc.save(docx_path)
    logger.info("DOCX 已创建: %s", docx_path)
    if src.is_dir():
        import shutil
        shutil.rmtree(src, ignore_errors=True)
    elif src.is_file():
        src.unlink(missing_ok=True)
    return docx_path


# ── JMComic Handler ──────────────────────────────────────────────

async def handle_jm(args: list[str], event: Event) -> str:
    if not args:
        return "用法: /jm <漫画ID> [输出目录]"

    jm_id = args[0]
    output_dir = args[1] if len(args) > 1 else str(JM_OUTPUT)

    await _send_reply(event, f"⏳ 正在下载 JM {jm_id}，请稍候...")

    try:
        result = await _get_mcp().download_jm_comic(jm_id=jm_id, output_dir=output_dir)
    except Exception as e:
        logger.error(f"JM download failed: {e}")
        return f"❌ JM 下载失败: {e}"

    if result.endswith(".pdf"):
        pdf_path = result
        try:
            await _send_file(event, pdf_path)
            return f"✅ JM {jm_id} 下载完成"
        except Exception as e:
            logger.error(f"JM file send failed: {e}")
            return f"⚠️ JM {jm_id} 已下载，但文件发送失败"
    return f"❌ {result}"


# ── Pixiv Handlers ───────────────────────────────────────────────

async def handle_pixiv_search(args: list[str], event: Event) -> str:
    if not args or args[0] in ("-h", "--help"):
        return ("用法: /pixiv search <关键词> [-t 搜索目标] [-s 排序] [-d 时段]\n"
                "搜索目标: tags(默认)/caption/keyword\n"
                "排序: date_desc(默认)/date_asc/popular_desc/popular_asc\n"
                "时段: day/week/month")

    word = args[0]
    kwargs = {"word": word}
    i = 1
    while i < len(args):
        if args[i] == "-t" and i + 1 < len(args):
            m = {"tags": "partial_match_for_tags",
                 "caption": "title_and_caption",
                 "keyword": "keyword"}
            kwargs["search_target"] = m.get(args[i + 1], "partial_match_for_tags")
            i += 2
        elif args[i] == "-s" and i + 1 < len(args):
            kwargs["sort"] = args[i + 1]
            i += 2
        elif args[i] == "-d" and i + 1 < len(args):
            kwargs["duration"] = f"within_last_{args[i + 1]}"
            i += 2
        elif args[i] == "-o" and i + 1 < len(args):
            try:
                kwargs["offset"] = int(args[i + 1])
            except ValueError:
                return "❌ -o 参数必须是数字"
            i += 2
        else:
            i += 1

    try:
        result = await _get_mcp().search_illust(**kwargs)
        data = json.loads(result)
    except Exception as e:
        logger.error("pixiv search error: %s", e)
        return f"❌ Pixiv 搜索失败: {e}"
    illusts = data.get("illusts", [])
    if not illusts:
        return "未找到结果"

    lines = [f'🔍 搜索 "{word}" 结果 ({len(illusts)} 条):']
    for il in illusts[:10]:
        lines.append(f"  [{il.get('id','?')}] {il.get('title','?')} — {il.get('user_name','?')} "
                     f"(♡{il.get('total_bookmarks',0)} 👁{il.get('total_view',0)})")
    if len(illusts) > 10:
        lines.append(f"  ... 还有 {len(illusts) - 10} 条")
    return "\n".join(lines)


async def handle_pixiv_rank(args: list[str], event: Event) -> str:
    mode = args[0] if args else "day"
    date = args[1] if len(args) > 1 else None
    kwargs = {"mode": mode}
    if date:
        kwargs["date"] = date

    try:
        result = await _get_mcp().illust_ranking(**kwargs)
        data = json.loads(result)
    except Exception as e:
        logger.error("pixiv rank error: %s", e)
        return f"❌ Pixiv 排行榜获取失败: {e}"
    illusts = data.get("illusts", [])

    if not illusts:
        return "排行榜暂无数据"

    lines = [f"🏆 Pixiv 排行榜 ({mode})"]
    for i, il in enumerate(illusts[:15], 1):
        lines.append(f"  {i}. [{il.get('id','?')}] {il.get('title','?')} — {il.get('user_name','?')}")
    return "\n".join(lines)


async def handle_pixiv_detail(args: list[str], event: Event) -> str:
    if not args:
        return "用法: /pixiv detail <作品ID>"
    try:
        result = await _get_mcp().illust_detail(illust_id=int(args[0]))
    except ValueError:
        return f"❌ 作品 ID 必须是数字: {args[0]}"
    except Exception as e:
        msg = str(e)
        if "404" in msg or "不存在" in msg or "not found" in msg.lower():
            return f"❌ 作品 {args[0]} 不存在"
        logger.error("pixiv detail error: %s", e)
        return f"❌ 查询作品详情失败: {e}"
    data = json.loads(result)
    if "error" in data:
        return f"❌ {data['error']}"

    tags = ", ".join(t.get("name", "") for t in data.get("tags", [])[:8])
    return (f"📖 {data.get('title','?')}\n"
            f"ID: {data.get('id','?')}  |  类型: {data.get('type','?')}\n"
            f"作者: {data.get('user_name','?')} ({data.get('user_id','?')})\n"
            f"页数: {data.get('page_count','?')}  |  尺寸: {data.get('width')}×{data.get('height')}\n"
            f"收藏: {data.get('total_bookmarks','?')}  |  浏览: {data.get('total_view','?')}\n"
            f"日期: {data.get('create_date','?')}\n"
            f"标签: {tags}")


async def handle_pixiv_related(args: list[str], event: Event) -> str:
    if not args:
        return "用法: /pixiv related <作品ID>"
    try:
        result = await _get_mcp().illust_related(illust_id=int(args[0]))
    except ValueError:
        return f"❌ 作品 ID 必须是数字: {args[0]}"
    except Exception as e:
        msg = str(e)
        if "404" in msg or "不存在" in msg or "not found" in msg.lower():
            return f"❌ 作品 {args[0]} 不存在"
        logger.error("pixiv related error: %s", e)
        return f"❌ 查询相关作品失败: {e}"
    data = json.loads(result)
    illusts = data.get("illusts", [])
    if not illusts:
        return "无相关作品"
    lines = [f"🔗 相关作品 ({len(illusts)} 条):"]
    for il in illusts[:10]:
        lines.append(f"  [{il.get('id','?')}] {il.get('title','?')} — {il.get('user_name','?')}")
    return "\n".join(lines)


async def handle_pixiv_recommend(args: list[str], event: Event) -> str:
    try:
        result = await _get_mcp().illust_recommended()
        data = json.loads(result)
    except Exception as e:
        logger.error("pixiv recommend error: %s", e)
        return f"❌ Pixiv 推荐获取失败: {e}"
    illusts = data.get("illusts", [])
    if not illusts:
        return "暂无推荐"
    lines = [f"🎯 推荐作品:"]
    for il in illusts[:10]:
        lines.append(f"  [{il.get('id','?')}] {il.get('title','?')} — {il.get('user_name','?')}")
    return "\n".join(lines)


async def handle_pixiv_trending(args: list[str], event: Event) -> str:
    try:
        result = await _get_mcp().trending_tags_illust()
        data = json.loads(result)
    except Exception as e:
        logger.error("pixiv trending error: %s", e)
        return f"❌ Pixiv 热门标签获取失败: {e}"
    tags = data.get("trending_tags", [])
    if not tags:
        return "暂无热门标签"
    lines = [f"🔥 热门标签:"]
    for t in tags[:15]:
        name = t.get("translated_name") or t.get("tag", "")
        lines.append(f"  #{name}")
    return "\n".join(lines)


async def handle_pixiv_download(args: list[str], event: Event) -> str:
    if not args:
        return "用法: /pixiv dl <作品ID> [输出目录]"
    try:
        illust_id = int(args[0])
    except ValueError:
        return f"❌ 作品 ID 必须是数字，你输入的是: {args[0]}"
    output_dir = args[1] if len(args) > 1 else str(PIXIV_OUTPUT)

    await _send_reply(event, f"⏳ 正在下载作品 {illust_id}，请稍候...")

    loop = asyncio.get_event_loop()

    try:
        result = await _get_mcp().download(illust_id=illust_id, output_dir=output_dir, wait=True)
        data = json.loads(result)
    except Exception as e:
        msg = str(e)
        if "404" in msg or "不存在" in msg:
            return f"❌ 作品 {illust_id} 不存在"
        return f"❌ 下载失败: {e}"

    if data.get("status") == "completed":
        result_path = data.get("result_path", "")
        if result_path:
            rp = str(result_path)
            try:
                if rp.endswith(".gif"):
                    docx_path = await loop.run_in_executor(None, _create_docx, rp)
                    await _send_file(event, docx_path)
                else:
                    pdf_path = await loop.run_in_executor(None, _create_pdf, rp)
                    await _send_file(event, pdf_path)
            except Exception as e:
                logger.error(f"Pixiv file send failed: {e}")
                return f"⚠️ 作品 {illust_id} 已下载，但文件发送失败"
        return "✅ 下载完成"
    elif data.get("error"):
        return f"❌ 下载失败: {data['error']}"
    return f"结果:\n{result[:500]}"


async def handle_pixiv_user(args: list[str], event: Event) -> str:
    if not args:
        return "用法: /pixiv user <用户ID>"
    try:
        result = await _get_mcp().search_user(word=args[0])
    except Exception as e:
        msg = str(e)
        if "OAuth" in msg or "认证" in msg:
            return "❌ 用户搜索需要 OAuth 认证，当前 PHPSESSID 模式不支持"
        logger.error("pixiv user search error: %s", e)
        return f"❌ 搜索用户失败: {e}"
    data = json.loads(result)
    users = data.get("users", [])
    if not users:
        return "未找到用户"
    u = users[0]
    return (f"👤 {u.get('name','?')} (@{u.get('account','?')})\n"
            f"ID: {u.get('id','?')}\n"
            f"简介: {u.get('comment', '')[:200]}")


async def handle_pixiv_bookmarks(args: list[str], event: Event) -> str:
    if not args:
        return "用法: /pixiv bookmarks <用户ID>"
    try:
        result = await _get_mcp().user_bookmarks(user_id=int(args[0]))
    except ValueError:
        return f"❌ 用户 ID 必须是数字: {args[0]}"
    except Exception as e:
        msg = str(e)
        if "OAuth" in msg or "认证" in msg:
            return "❌ 用户收藏需要 OAuth 认证，当前 PHPSESSID 模式不支持"
        if "404" in msg or "不存在" in msg:
            return f"❌ 用户 {args[0]} 不存在"
        logger.error("pixiv bookmarks error: %s", e)
        return f"❌ 查询用户收藏失败: {e}"
    data = json.loads(result)
    illusts = data.get("illusts", [])
    if not illusts:
        return "该用户无公开收藏"
    lines = [f"📑 用户 {args[0]} 的收藏 ({len(illusts)} 条):"]
    for il in illusts[:10]:
        lines.append(f"  [{il.get('id','?')}] {il.get('title','?')}")
    return "\n".join(lines)


# ── Pixiv Router ─────────────────────────────────────────────────

PIXIV_HANDLERS = {
    "search": handle_pixiv_search,
    "rank": handle_pixiv_rank,
    "detail": handle_pixiv_detail,
    "related": handle_pixiv_related,
    "recommend": handle_pixiv_recommend,
    "trending": handle_pixiv_trending,
    "dl": handle_pixiv_download,
    "download": handle_pixiv_download,
    "user": handle_pixiv_user,
    "bookmarks": handle_pixiv_bookmarks,
}


async def handle_pixiv(args: list[str], event: Event) -> str:
    if not args or args[0] in ("-h", "--help"):
        return ("Pixiv 命令:\n"
                "  /pixiv search <词>     搜索插画\n"
                "  /pixiv rank [模式]     排行榜\n"
                "  /pixiv detail <ID>     作品详情\n"
                "  /pixiv related <ID>    相关作品\n"
                "  /pixiv dl <ID>         下载作品\n"
                "  /pixiv recommend       推荐\n"
                "  /pixiv trending        热门标签\n"
                "  /pixiv user <ID>       用户信息\n"
                "  /pixiv bookmarks <ID>  用户收藏\n"
                "\n搜索高级选项:\n"
                "  /pixiv search <词> -t tags|caption|keyword -s sort -d day|week|month")

    sub = args[0]
    handler = PIXIV_HANDLERS.get(sub)
    if not handler:
        return f"未知子命令: {sub}。使用 /pixiv -h 查看帮助。"
    return await handler(args[1:], event)


# ── Bilibili Handler ────────────────────────────────────────────

async def handle_bili(args: list[str], event: Event) -> str:
    if not args:
        return "用法: /bili <B站视频链接 或 BV号>"

    url = args[0]
    if not url.startswith("http"):
        bv = extract_bvid(url)
        url = f"https://www.bilibili.com/video/{bv}" if bv else url

    await _send_reply(event, "⏳ 正在下载 B站视频，请稍候...")

    try:
        result = await _get_mcp().bili_download_video(url=url, output_dir=str(BILI_OUTPUT))
    except Exception as e:
        logger.error(f"Bilibili download failed: {e}")
        return f"❌ B站下载失败: {e}"

    if not result or "失败" in result or "错误" in result or "不存在" in result:
        return f"❌ {result}"

    try:
        await _send_video(event, result)
    except Exception as e:
        logger.error(f"Bilibili video send failed: {e}")
        return f"⚠️ 视频已下载但发送失败: {e}"

    display = extract_bvid(args[0]) or args[0].split("/")[-1]
    return f"✅ B站视频 {display} 下载完成"


# ── Live 分页会话 ────────────────────────────────────────────────

LIVE_PAGE_SIZE = 10
LIVE_CACHE_TTL = 600
_live_sessions: dict[str, dict] = {}


def _session_key(event: Event) -> str:
    if event.message_type == "group":
        return f"group:{event.group_id}"
    return f"private:{event.user_id}"


def _get_live_session(event: Event) -> dict | None:
    key = _session_key(event)
    sess = _live_sessions.get(key)
    if not sess:
        return None
    if time.time() - sess.get("ts", 0) > LIVE_CACHE_TTL:
        _live_sessions.pop(key, None)
        return None
    return sess


def _set_live_session(event: Event, items: list[dict], page: int = 0) -> dict:
    sess = {"items": items, "page": page, "ts": time.time()}
    _live_sessions[_session_key(event)] = sess
    return sess


def _clear_live_session(event: Event) -> None:
    _live_sessions.pop(_session_key(event), None)


async def handle_live(args: list[str], event: Event) -> str:
    sess = _get_live_session(event)
    if sess:
        sess["ts"] = time.time()
        return format_live_page(sess["items"], sess["page"], LIVE_PAGE_SIZE)

    await _send_reply(event, "⏳ 正在查询你关注的主播直播情况，请稍候...")
    try:
        raw = await _get_mcp().bili_live_following()
        data = json.loads(raw)
    except Exception as e:
        logger.error("bili_live_following error: %s", e)
        return f"❌ 查询直播状态失败: {e}"

    if data.get("error"):
        err = data["error"]
        if "cookies" in err and "缺少" in err:
            return f"❌ {err}"
        elif "失效" in err:
            return f"❌ {err}"
        elif "消息" in data:
            return f"📋 {data['message']}"
        return f"❌ {err}"

    items = data.get("live", [])
    if not items:
        return "😴 你关注的主播当前没有人在直播"
    sess = _set_live_session(event, items, 0)
    return format_live_page(sess["items"], 0, LIVE_PAGE_SIZE)


async def handle_live_nav(action: str, args: list[str], event: Event) -> str:
    sess = _get_live_session(event)

    if action == "break":
        if sess is None:
            return "没有正在进行的直播浏览，先用 /live 开始查询吧"
        _clear_live_session(event)
        return "👋 已退出直播浏览"

    if sess is None:
        return "缓存已过期或未开始查询，请先发 /live"

    total = len(sess["items"])
    total_pages = (total + LIVE_PAGE_SIZE - 1) // LIVE_PAGE_SIZE
    cur = sess["page"]

    if action == "next":
        if cur >= total_pages - 1:
            return f"已是最后一页 (第 {cur + 1}/{total_pages} 页)"
        cur += 1
    elif action == "prev":
        if cur <= 0:
            return f"已是第一页 (第 1/{total_pages} 页)"
        cur -= 1
    elif action == "goto":
        if not args or not args[0].isdigit():
            return "用法: /第N页 <数字>  例如 /第3页"
        n = int(args[0])
        if n < 1:
            n = 1
        if n > total_pages:
            n = total_pages
        cur = n - 1
        if cur == sess["page"]:
            return f"已经在第 {n}/{total_pages} 页"

    sess["page"] = cur
    sess["ts"] = time.time()
    return format_live_page(sess["items"], cur, LIVE_PAGE_SIZE)


# ── Message Handler ──────────────────────────────────────────────

@bot.on_message()
async def handle_message(event: Event) -> None:
    raw = event.raw_message.strip()
    logger.info("收到消息: type=%s user=%s text=%r", event.message_type, event.user_id, raw)
    if not raw:
        return

    try:
        await _dispatch_message(event, raw)
    except Exception as e:
        logger.error("handle_message 顶层异常", exc_info=True)
        try:
            await _send_reply(event, f"❌ 处理消息时出错: {e}\n输入 /help 查看可用命令。")
        except Exception:
            pass


async def _dispatch_message(event: Event, raw: str) -> None:
    text = raw
    is_group = event.message_type == "group"
    user_id = event.user_id

    if is_group:
        at_match = re.match(r'^\[CQ:at,qq=(\d+)\]\s*(.*)', raw)
        if at_match:
            at_qq = at_match.group(1)
            if at_qq not in ("2981572633", os.environ.get("BOT_QQ", "")):
                return
            text = at_match.group(2).strip()
            if not text:
                return
        elif not raw.startswith("/"):
            return

    if text in ("/help", "help", "/start") or text.startswith("帮助"):
        await _send_reply(event, (
            "🤖 QQ Bot 命令:\n"
            "  /jm <ID>          下载 JM 漫画 → PDF\n"
            "  /pixiv <命令>      Pixiv 搜索/浏览/下载\n"
            "  /bili <链接>       下载 B站视频\n"
            "  /live             查看我关注的主播谁在直播（分页浏览）\n"
            "    /下一页 /上一页 /第N页 /break 翻页或退出\n"
            "  /搜直播 <关键词>   搜索 B站直播（如 /搜直播 原神）\n"
            "  /news             游戏新闻菜单（国内 / 国外）\n"
            "    /国内新闻        ali213 最新 5 条\n"
            "    /国外新闻        显示国外源子菜单（VGC/ResetEra/Insider/Steam）\n"
            "  /help             显示此帮助\n"
            "  /forget           清空 AI 对话记忆\n"
            "\n💬 群聊中 @我 可以和我对话，AI 会自动处理你的请求。\n"
            "例如：「帮我搜一下碧蓝档案的图」「今天P站排行榜」"
        ))
        return

    jm_match = re.match(r"^/jm(?:\s+(.+))?$", text)
    if jm_match:
        args = _split_args(jm_match.group(1) or "")
        result = await handle_jm(args, event)
        await _send_reply(event, result)
        return

    pixiv_match = re.match(r"^/pixiv(?:\s+(.+))?$", text)
    if pixiv_match:
        args = _split_args(pixiv_match.group(1) or "")
        result = await handle_pixiv(args, event)
        await _send_reply(event, result)
        return

    if text == "/news":
        await _send_reply(event, (
            "📰 游戏新闻菜单\n"
            "  /国内新闻  — ali213 最新 5 条（带图）\n"
            "  /国外新闻  — 显示国外源子菜单"
        ))
        return

    if text == "/国内新闻":
        await _send_reply(event, "📡 正在抓取国内新闻...")
        loop = asyncio.get_event_loop()
        try:
            result = await loop.run_in_executor(None, fetch_news)
            await _send_reply(event, result)
        except Exception as e:
            logger.error("国内新闻抓取失败: %s", e)
            await _send_reply(event, f"❌ 国内新闻抓取失败: {e}")
        return

    if text == "/国外新闻":
        await _send_reply(event, (
            "🌍 国外游戏热料源（各取最新 5 条，带图+中文翻译）\n"
            "  /vgc      — VGC 新闻\n"
            "  /resetera — ResetEra 爆料帖\n"
            "  /insider  — Insider Gaming（含 Bluesky 镜像）\n"
            "  /steam    — Steam 即将发售\n"
            "💡 翻译需 5-15 秒，请稍候"
        ))
        return

    src_match = re.match(r"^/(vgc|resetera|insider|steam)$", text)
    if src_match:
        name = src_match.group(1)
        fn = SOURCE_COMMANDS.get(name)
        if fn:
            await _send_reply(event, f"🌍 正在抓取 {name}，需翻译请稍候...")
            loop = asyncio.get_event_loop()
            try:
                result = await loop.run_in_executor(None, fn)
                await _send_reply(event, result)
            except Exception as e:
                logger.error("%s 抓取失败: %s", name, e)
                await _send_reply(event, f"❌ {name} 抓取失败: {e}")
            return

    if text in ("/live", "/直播"):
        result = await handle_live([], event)
        await _send_reply(event, result)
        return

    search_live_match = re.match(r"^/(?:搜直播|slive)\s+(.+)", text)
    if search_live_match:
        keyword = search_live_match.group(1).strip()
        await _send_reply(event, f"🔍 正在搜索直播「{keyword}」...")
        try:
            raw = await _get_mcp().bili_search_live(keyword=keyword)
            data = json.loads(raw)
        except Exception as e:
            await _send_reply(event, f"❌ 搜索失败: {e}")
            return
        if data.get("error"):
            await _send_reply(event, f"❌ {data['error']}")
            return
        rooms = data.get("rooms", [])
        if not rooms:
            await _send_reply(event, f"未找到「{keyword}」相关直播")
            return
        lines = [f"🔍 搜索「{keyword}」直播结果:"]
        for r in rooms[:10]:
            face = f"[CQ:image,file={r['face']}]" if r.get("face") else ""
            cover = f"[CQ:image,file={r['cover']}]" if r.get("cover") else ""
            lines.append(
                f"{face} {r['name']}\n"
                f"📺 {r['title'][:40]}\n"
                f"{cover}\n"
                f"🏷 {r.get('area','')}  👥{r['online']}  {r['link']}"
            )
        await _send_reply(event, "\n".join(lines))
        return

    if text in ("/下一页", "/next", "/下页"):
        result = await handle_live_nav("next", [], event)
        await _send_reply(event, result)
        return
    if text in ("/上一页", "/prev", "/上页"):
        result = await handle_live_nav("prev", [], event)
        await _send_reply(event, result)
        return
    if text in ("/break", "/退出"):
        result = await handle_live_nav("break", [], event)
        await _send_reply(event, result)
        return
    goto_match = re.match(r"^/(?:第\s*(\d+)\s*页|page)\s*(\d+)?$", text)
    if goto_match:
        n = goto_match.group(1) or goto_match.group(2)
        if not n:
            await _send_reply(event, "用法: /第N页 <数字>  例如 /第3页")
            return
        result = await handle_live_nav("goto", [n], event)
        await _send_reply(event, result)
        return

    bili_match = re.match(r"^/bili(?:de)?(?:\s+(.+))?$", text)
    if bili_match:
        args = _split_args(bili_match.group(1) or "")
        result = await handle_bili(args, event)
        await _send_reply(event, result)
        return

    if text == "/forget":
        try:
            ai = _get_ai()
            cid = f"group:{event.group_id}" if is_group else f"private:{event.user_id}"
            ai.clear_history(cid)
            await _send_reply(event, "🧹 已清空对话记忆，我们重新开始吧～")
        except Exception:
            await _send_reply(event, "🧹 已清空对话记忆")
        return

    try:
        ai = _get_ai()
        cid = f"group:{event.group_id}" if is_group else f"private:{event.user_id}"
        logger.info("Sending to AI (cid=%s): %s", cid, text)

        async def _on_tool_start(tool_name: str):
            if tool_name.startswith("download"):
                await _send_reply(event, f"⏳ 正在{tool_name.replace('_', ' ')}，请稍候...")

        reply, files = await ai.handle(text, user_id=cid, on_tool_start=_on_tool_start)
        logger.info("AI reply (first 200): %s", reply[:200] if reply else "None")
        await _send_reply(event, reply)
        for f in files:
            if f.lower().endswith((".mp4", ".mkv", ".webm", ".flv", ".avi")):
                await _send_video(event, f)
            else:
                await _send_file(event, f)
        logger.info("AI reply sent, files=%d", len(files))
    except ValueError as e:
        logger.warning("AI not available: %s", e)
        if "OPENAI_API_KEY" in str(e):
            await _send_reply(event, "❌ AI 功能未配置，请在 .env 中设置 OPENAI_API_KEY")
        else:
            await _send_reply(event, f"❌ AI 初始化失败: {e}")
        logger.info("Fallback reply sent")
    except Exception as e:
        logger.error("AI handler error", exc_info=True)
        await _send_reply(event, f"抱歉，处理出错: {e}\n输入 /help 查看可用命令。")
        logger.info("Error reply sent")


# ── Main ─────────────────────────────────────────────────────────

@bot.server_app.before_serving
async def _startup_mcp():
    global _mcp_manager, _ai_handler
    logger.info("启动 MCP 服务...")
    _mcp_manager = McpManager()
    await _mcp_manager.start()
    _ai_handler = AIHandler(_mcp_manager)
    logger.info("MCP 服务启动完成")


def main():
    host = os.environ.get("BOT_HOST", "127.0.0.1")
    port = int(os.environ.get("BOT_PORT", "8080"))

    logger.info("=" * 50)
    logger.info("QQ Bot 启动")
    logger.info("WebSocket 服务: %s:%s", host, port)
    logger.info("NapCatQQ 配置: ws://%s:%s/ws/", host, port)
    logger.info("OpenAI API Key: %s", "已设置" if os.environ.get("OPENAI_API_KEY") else "未设置")
    logger.info("OpenAI Model: %s", os.environ.get("OPENAI_MODEL", "gpt-4o-mini"))
    logger.info("OpenAI Base URL: %s", os.environ.get("OPENAI_BASE_URL", "https://api.openai.com/v1"))
    logger.info("=" * 50)
    logger.info("Bot started. Waiting for messages...")

    bot.run(host=host, port=port)


if __name__ == "__main__":
    main()
